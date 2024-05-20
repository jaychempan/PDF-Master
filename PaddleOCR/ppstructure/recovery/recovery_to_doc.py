# Copyright (c) 2020 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
from copy import deepcopy

from ppstructure.recovery.table_process import HtmlToDocx

from ppocr.utils.logging import get_logger

logger = get_logger()

def convert_info_to_json(img, res, save_folder, img_name):
    import json
    pdf_info = {
        "pdf_info": [],
        "_parse_type": "ocr"
    }
    page_idx = -1
    for page in res:
        page_idx = page_idx + 1
        # 这里双栏pdf的话，就是把双栏变成单栏为一页
        page_info = {
                "preproc_blocks": [],
                # "layout_bboxes": [], # 感觉不需要
                "page_idx": page_idx, # 需要更新
                # "page_size": [595.0, 842.0], # 不重要暂时不处理
                # "_layout_tree": [], # 好像和layout_bboxes没区别
                "images": [], # 图像包括图像标题放这里
                "tables": [], # 表格包括表格标题放这里
                "interline_equations": [],
                "discarded_blocks": [], # 需要丢弃的放这里
                "para_blocks": [] # 按照段落顺序，但是好像没区别
            }
        
        for region in page:
            if len(region["res"]) == 0:
                continue
            # 处理图片和图片标题，放到images里面
            if region["type"].lower() == 'figure':
                block = {
                "type": region["type"].lower(),
                "bbox": region["bbox"],
                "lines": []
                }
                for span in region["res"]:
                    block["lines"].append({
                        "bbox": region["bbox"],
                        "spans": [
                            {
                                "bbox": region["bbox"],
                                "type": "image_body", # 是文字的话就是text，但是如果是inline_equation的话就需要进行格式化转换（待做）
                                "image_path": "./figures/{}_{}.jpg".format(page_idx, region["bbox"])
                            }
                        ]
                    })
                page_info["images"].append(block)
            elif region["type"].lower() == 'figure_caption':
                block = {
                "type": region["type"].lower(),
                "bbox": region["bbox"],
                "lines": []
                }
                for span in region["res"]:
                    block["lines"].append({
                        "bbox": region["bbox"],
                        "spans": [
                            {
                                "bbox": region["bbox"],
                                "content": span["text"],
                                "type": "image_caption", # 是文字的话就是text，但是如果是inline_equation的话就需要进行格式化转换（待做）
                                
                            }
                        ]
                    })
                page_info["images"].append(block)
            elif region["type"].lower() == 'table':
                block = {
                "type": region["type"].lower(),
                "bbox": region["bbox"],
                "lines": []
                }
                for span in region["res"]:
                    block["lines"].append({
                        "bbox": region["bbox"],
                        "spans": [
                            {
                                "bbox": region["bbox"],
                                "type": "table", # 是文字的话就是text，但是如果是inline_equation的话就需要进行格式化转换（待做）
                                "image_path": "./tables/{}_{}.jpg".format(page_idx, region["bbox"])
                            }
                        ]
                    })
                page_info["tables"].append(block)
            elif region["type"].lower() == 'table_caption':
                block = {
                "type": region["type"].lower(),
                "bbox": region["bbox"],
                "lines": []
                }
                for span in region["res"]:
                    block["lines"].append({
                        "bbox": region["bbox"],
                        "spans": [
                            {
                                "bbox": region["bbox"],
                                "content": span["text"],
                                "type": "table_caption", # 是文字的话就是text，但是如果是inline_equation的话就需要进行格式化转换（待做）
                            }
                        ]
                    })
                page_info["tables"].append(block)
            elif region["type"].lower() == 'title':
                block = {
                "type": region["type"].lower(),
                "bbox": region["bbox"],
                "lines": []
                }
                for span in region["res"]:
                    block["lines"].append({
                        "bbox": region["bbox"],
                        "spans": [
                            {
                                "bbox": span["text_region"],
                                "content": span["text"],
                                "type": "text" # 是文字的话就是text，但是如果是inline_equation的话就需要进行格式化转换（待做）
                            }
                        ]
                    })
                page_info["preproc_blocks"].append(block)

            elif region["type"].lower() == 'text':
                block = {
                "type": region["type"].lower(),
                "bbox": region["bbox"],
                "lines": []
                }
                for span in region["res"]:
                    block["lines"].append({
                        "bbox": region["bbox"],
                        "spans": [
                            {
                                "bbox": region["bbox"],
                                "content": span["text"],
                                "type": "text" # 是文字的话就是text，但是如果是inline_equation的话就需要进行格式化转换（待做）
                            }
                        ]
                    })
                page_info["preproc_blocks"].append(block)

            elif region["type"].lower() == 'equation':
                block = {
                "type": region["type"].lower(),
                "bbox": region["bbox"],
                "lines": []
                }
                for span in region["res"]:
                    block["lines"].append({
                        "bbox": region["bbox"],
                        "spans": [
                            {
                                "bbox": region["bbox"],
                                "content": span["text"],
                                "type": "text",
                                "image_path": "./equations/{}_{}.jpg".format(page_idx, region["bbox"])
                            }
                        ]
                    })
            page_info["preproc_blocks"].append(block)

        pdf_info["pdf_info"].append(page_info)



    # Save to JSON file
    json_path = os.path.join(save_folder, f"{img_name}_ocr.json")
    with open(json_path, 'w', encoding='utf-8') as json_file:
        json.dump(pdf_info, json_file, ensure_ascii=False, indent=4)
    print(f"JSON saved to {json_path}")



# 转换成markdown格式
def convert_info_markdown(img, res, save_folder, img_name):
    import markdownify
    markdown_content = []

    for region in res:
        if len(region["res"]) == 0:
            continue

        if region["type"].lower() == "figure":
            img_idx = region["img_idx"]
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = "./figures/{}_{}.jpg".format(img_idx, region["bbox"])
            markdown_content.append(f"![Figure {img_idx}]({img_path})\n")
        
        elif region["type"].lower() == "equation":
            # Ignored as specified in the original code
            img_idx = region["img_idx"]
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = "./equations/{}_{}.jpg".format(img_idx, region["bbox"])
            markdown_content.append(f"![Equations {img_idx}]({img_path})\n")

        elif region["type"].lower() == "title":
            level = "#"  # Adjust as needed for different levels of headings
            markdown_content.append(f"{level} {region['res'][0]['text']}\n")
        
        elif region["type"].lower() == "table":
            html_content = region["res"]["html"]
            markdown_table = markdownify.markdownify(html_content, heading_style="ATX")
            markdown_content.append(markdown_table + "\n")

        # 需要对text进行合并操作，因为前面已经进行了sorted_layout_boxes操作，所以这里的文字可以直接进行合并
        elif region["type"].lower() == "text":
            spans = ""
            for line in region["res"]:
                if line['text'].strip().endswith('.') or line['text'].strip().endswith('。'):
                    # print(line['text'])
                    # 提取行末尾前的单词
                    words = line['text'].rsplit('.', 1)[0].lower().split()
                    # print(words)
                    if words[-1] in ['fig', 'tab', '图', '表']:
                        spans += line['text'] + ' '
                    else:
                        spans += line['text'] + '\n\n'
                elif line['text'].strip().endswith('-'):
                    spans += line['text'].rstrip('-')
                else:
                    spans += line['text'] + ' '
            markdown_content.append(f"{spans}")

        # 过滤无用的图片标题表格标题（前面已经删除了页眉和页脚）
        elif region["type"].lower() in ["figure_caption", "table_caption"]:
            # Ignored as specified in the original code
            # continue
            pass
        else:
            for line in region["res"]:
                markdown_content.append(f"{line['text']}\n")
    
    # Save to markdown file
    markdown_path = os.path.join(save_folder, f"{img_name}_ocr.md")
    with open(markdown_path, 'w', encoding='utf-8') as md_file:
        md_file.write("\n".join(markdown_content))
    print(f"Markdown saved to {markdown_path}")


# 转换成docx格式
def convert_info_docx(img, res, save_folder, img_name):
    from docx import Document
    from docx import shared
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = shared.Pt(6.5)

    flag = 1
    for i, region in enumerate(res):
        if len(region["res"]) == 0:
            continue
        img_idx = region["img_idx"]
        if flag == 2 and region["layout"] == "single":
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "1")
            flag = 1
        elif flag == 1 and region["layout"] == "double":
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "2")
            flag = 2

        if region["type"].lower() == "figure":
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(
                # excel_save_folder, "{}_{}.jpg".format(region["bbox"], img_idx)
                excel_save_folder, "figures/{}_{}.jpg".format(img_idx, region["bbox"])
            )
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif region["type"].lower() == "title":
            doc.add_heading(region["res"][0]["text"])
        elif region["type"].lower() == "table":
            parser = HtmlToDocx()
            parser.table_style = "TableGrid"
            parser.handle_table(region["res"]["html"], doc)
        # 对页眉页脚不做处理，即丢弃
        elif region["type"].lower() == "header" or region["type"].lower() == "footer":
            pass
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for i, line in enumerate(region["res"]):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line["text"] + " ")
                text_run.font.size = shared.Pt(10)

    # save to docx
    docx_path = os.path.join(save_folder, "{}_ocr.docx".format(img_name))
    doc.save(docx_path)
    logger.info("docx save to {}".format(docx_path))


def sorted_layout_boxes(res, w):
    # 根据布局来识别顺序，可以加上页内不同板块中文字段落的识别,将不同行的进行合并，但是段落间又分开
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]["layout"] = "single"
        return res

    sorted_boxes = sorted(res, key=lambda x: (x["bbox"][1], x["bbox"][0]))
    _boxes = list(sorted_boxes)
    
    new_res = [] # 存放最终的排好序的结果
    res_left = [] # 如果有左边结果存放到这里面
    res_right = [] # 如果有右边的结果存档到这里，最终在new_res中进行排序
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if (
                _boxes[i]["bbox"][1] > _boxes[i - 1]["bbox"][3]
                and _boxes[i]["bbox"][0] < w / 2
                and _boxes[i]["bbox"][2] > w / 2
            ):
                new_res += res_left
                new_res += res_right
                _boxes[i]["layout"] = "single"
                new_res.append(_boxes[i])
            else:
                if _boxes[i]["bbox"][2] > w / 2:
                    _boxes[i]["layout"] = "double"
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]["bbox"][0] < w / 2:
                    _boxes[i]["layout"] = "double"
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]["bbox"][0] < w / 4 and _boxes[i]["bbox"][2] < 3 * w / 4:
            _boxes[i]["layout"] = "double"
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]["bbox"][0] > w / 4 and _boxes[i]["bbox"][2] > w / 2:
            _boxes[i]["layout"] = "double"
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]["layout"] = "single"
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res

